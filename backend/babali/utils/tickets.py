import copy
import docx
import os
import subprocess

from threading import Timer
from django.conf import settings


TEMPLATES_ROOT = 'static/tickets/templates/'
TICKETS_ROOT = 'static/tickets/'
TICKETS_PDF_DELETE_INTERVAL = 30

def _replace_text_in_paragraph(paragraph, data, placeholder_map):
    for placeholder, key in placeholder_map.items():
        if placeholder in paragraph.text:
            value = data.get(key, '')
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)


def _fill_template(doc, data, placeholder_map):
    for p in doc.paragraphs:
        _replace_text_in_paragraph(p, data, placeholder_map)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_in_paragraph(p, data, placeholder_map)


def delete_ticket_pdf(file_path):
    if os.path.exists(file_path) and os.path.isfile(file_path):
        os.remove(file_path)

timer = Timer(TICKETS_PDF_DELETE_INTERVAL, delete_ticket_pdf, args=('.pdf', ))

def generate_tickets_pdf(template_name, ticket_type, output_name, tickets_data, placeholder_map):
    """
    Generates a single pdf document containing multiple tickets from a template.

    Args:
        template_path (str): The path to the .docx template file.
        output_path (str): The path for the combined output .docx file.
        tickets_data (list): A list of dictionaries, one for each ticket.
        placeholder_map (dict): A dictionary that maps placeholders in the
                                template to keys in the data dictionaries.
    """
    templates_dir = os.path.join(settings.BASE_DIR, TEMPLATES_ROOT)
    template_path = os.path.join(templates_dir, template_name)

    output_dir = os.path.join(settings.BASE_DIR, TICKETS_ROOT, ticket_type)
    output_doc_path = os.path.join(output_dir, f'{output_name}.docx')

    if not os.path.exists(template_path):
        raise Exception(f'error: Template file not found at "{template_path}"')
    if not isinstance(tickets_data, list) or not tickets_data:
        raise Exception('Error: tickets_data must be a non-empty list of dictionaries.')

    final_doc = docx.Document()

    for data in tickets_data:
        ticket_doc = docx.Document(template_path)
        data = {key: (str(value) if value is not None else '') for key, value in data.items()}
        print(data)
        _fill_template(ticket_doc, data, placeholder_map)

        for element in ticket_doc.element.body:
            final_doc.element.body.append(copy.deepcopy(element))

    try:
        final_doc.save(output_doc_path)
        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            output_doc_path
        ], check=True)
        os.remove(output_doc_path)

        output_pdf_path = os.path.join(output_dir, f'{output_name}.pdf')
        global timer
        timer.cancel()
        timer = Timer(TICKETS_PDF_DELETE_INTERVAL, delete_ticket_pdf, args=(output_pdf_path, ))
        timer.start()

        return output_pdf_path

    except Exception as e:
        print(f'error: {e}')